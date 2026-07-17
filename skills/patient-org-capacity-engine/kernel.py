"""Kernel helpers for patient-org-capacity-engine.

Threads one org_capacity_profile.json through a user-led, non-prescriptive
capacity-building engagement. See SKILL.md for the workflow.
"""
import json
import os
import datetime

SCHEMA_VERSION = "1.0"

BRANCHES = {
    "A_landscape": "Disease landscape & research-gap mapping",
    "B_readiness": "Research readiness & infrastructure (registry, biobank, natural history, models)",
    "C_pipeline": "Therapeutic strategy & pipeline (incl. CRISPR-led / genetic medicine)",
    "D_patientvoice": "Regulatory & patient voice (PFDD / EL-PFDD, Voice of the Patient)",
    "E_access": "Access & care navigation (diagnosis, expanded access, centers of care, payer/HTA)",
    "F_capacity": "Organizational capacity & funding (gap map, peer models, options by cost/horizon)",
}

MATURITY_TIERS = [
    "No tractable therapeutic strategy identified yet",
    "Building disease understanding & community",
    "Early research infrastructure (registry / natural history forming)",
    "Research-ready (assets exist; supporting external programs)",
    "Supporting or co-driving a therapeutic program",
    "Leading a therapeutic program in-house",
]

DISCOVERY_FIELDS = [
    "vision", "mission", "disease_area", "communities_served", "maturity",
    "budget_band", "staffing", "filing_type", "key_personnel",
    "current_assets", "risk_appetite", "constraints", "wants",
]


def now_iso():
    return datetime.datetime.now(datetime.timezone.utc).isoformat()


def init_profile(org_name, disease_area, path="org_capacity_profile.json"):
    """Create the source-of-truth profile and write it to disk. Returns dict."""
    profile = {
        "schema_version": SCHEMA_VERSION,
        "created_at": now_iso(),
        "org_name": org_name,
        "disease_area": disease_area,
        "discovery": {},
        "branches_selected": [],
        "stages": {},
        "assumptions": [],
        "checkpoints": [],
    }
    save_profile(profile, path)
    return profile


def load_profile(path="org_capacity_profile.json"):
    with open(path) as fh:
        return json.load(fh)


def save_profile(profile, path="org_capacity_profile.json"):
    profile["updated_at"] = now_iso()
    with open(path, "w") as fh:
        json.dump(profile, fh, indent=2)
    return path


def record_discovery(profile, path="org_capacity_profile.json", **fields):
    """Store discovery-interview answers. Unknown keys are kept but flagged."""
    unknown = [k for k in fields if k not in DISCOVERY_FIELDS]
    profile.setdefault("discovery", {}).update(fields)
    if unknown:
        profile["discovery"]["_extra_fields"] = sorted(
            set(profile["discovery"].get("_extra_fields", []) + unknown))
    save_profile(profile, path)
    return profile["discovery"]


def select_branches(profile, branch_ids, path="org_capacity_profile.json"):
    """Record the leader's chosen focus branches (validated against BRANCHES)."""
    if isinstance(branch_ids, str):
        branch_ids = [branch_ids]
    bad = [b for b in branch_ids if b not in BRANCHES]
    if bad:
        raise ValueError("Unknown branch id(s): %s. Valid: %s"
                         % (bad, list(BRANCHES)))
    profile["branches_selected"] = list(branch_ids)
    save_profile(profile, path)
    return [(b, BRANCHES[b]) for b in branch_ids]


def log_assumption(profile, statement, basis, how_to_resolve, stage=None,
                   path="org_capacity_profile.json"):
    """Append an assumption to the ledger. Do this every time you fill a blank.

    statement       - the assumption in one plain sentence
    basis           - why it was assumed (what was missing)
    how_to_resolve  - the single thing the leader could share to retire it
    """
    entry = {
        "index": len(profile.setdefault("assumptions", [])),
        "statement": statement,
        "basis": basis,
        "how_to_resolve": how_to_resolve,
        "stage": stage,
        "status": "open",
        "logged_at": now_iso(),
    }
    profile["assumptions"].append(entry)
    save_profile(profile, path)
    return entry


def resolve_assumption(profile, index, resolution,
                       path="org_capacity_profile.json"):
    """Mark an assumption resolved with the leader's real input."""
    for a in profile.get("assumptions", []):
        if a["index"] == index:
            a["status"] = "resolved"
            a["resolution"] = resolution
            a["resolved_at"] = now_iso()
            save_profile(profile, path)
            return a
    raise KeyError("No assumption with index %s" % index)


def pending_assumptions(profile):
    """The still-open assumptions to surface with the leader now."""
    return [a for a in profile.get("assumptions", []) if a["status"] == "open"]


def set_checkpoint(profile, stage, question, options=None,
                   path="org_capacity_profile.json"):
    """Record that you paused here for the leader. Pair with an ask_user call."""
    entry = {
        "stage": stage,
        "question": question,
        "options": options or [],
        "at": now_iso(),
    }
    profile.setdefault("checkpoints", []).append(entry)
    save_profile(profile, path)
    return entry


def update_stage(profile, stage, status, outputs=None, notes=None,
                 path="org_capacity_profile.json"):
    """Track progress and deliverables for a stage."""
    profile.setdefault("stages", {})[stage] = {
        "status": status,
        "outputs": outputs or [],
        "notes": notes,
        "at": now_iso(),
    }
    save_profile(profile, path)
    return profile["stages"][stage]


def assumptions_markdown(profile, include_resolved=False):
    """Render the ledger as a Markdown block for the top of a deliverable."""
    items = profile.get("assumptions", [])
    if not include_resolved:
        items = [a for a in items if a["status"] == "open"]
    lines = ["> **Assumptions this document rests on.** These are our best "
             "reading of gaps in what we were told; please correct any of them "
             "— each carries the one thing you could share to replace it.", ""]
    if not items:
        lines.append("> _No open assumptions — the inputs below were confirmed "
                     "with the organization._")
        return "\n".join(lines)
    for a in items:
        tag = "" if a["status"] == "open" else " _(resolved)_"
        lines.append("> - **%s**%s  \n>   _Why assumed:_ %s  \n>   _To resolve, "
                     "share:_ %s" % (a["statement"], tag, a["basis"],
                                     a["how_to_resolve"]))
    return "\n".join(lines)


def build_manifest(profile, path="capacity_manifest.json"):
    """Collect stages, outputs, branches, and open assumptions into a manifest."""
    manifest = {
        "org_name": profile.get("org_name"),
        "disease_area": profile.get("disease_area"),
        "branches": [(b, BRANCHES.get(b, b))
                     for b in profile.get("branches_selected", [])],
        "stages": profile.get("stages", {}),
        "deliverables": sorted({o for s in profile.get("stages", {}).values()
                                for o in s.get("outputs", [])}),
        "open_assumptions": pending_assumptions(profile),
        "n_checkpoints": len(profile.get("checkpoints", [])),
        "built_at": now_iso(),
    }
    with open(path, "w") as fh:
        json.dump(manifest, fh, indent=2)


# ---------------------------------------------------------------------------
# Deliverable rendering (added): editable Word (.docx) documents, a capacity
# -> strategy dependency model, and two roadmap views the leader can choose
# between. Requires python-docx and matplotlib (install into the env if
# missing). All third-party imports are deferred into function bodies.
# ---------------------------------------------------------------------------

# House palette as hex strings; RGBColor objects are built inside functions.
CLR_NAVY = "22223B"
CLR_TEAL = "2A9D8F"
CLR_GREY = "555555"
CLR_AMBER_BG = "FFF3E0"
POSTURE_COLORS = {
    "Amplify": "1982C4", "Connect": "2A9D8F",
    "Build-via-partner": "8A5A44", "Revive": "6A4C93",
}
NODE_KIND_COLORS = {
    "capacity": "6A4C93", "infrastructure": "1982C4", "science": "2A9D8F",
    "therapeutic": "E63946", "advocacy": "F4A261", "funding": "E9C46A",
}
HORIZON_TO_PHASE = {"H1": "0-12 months", "H2": "1-2 years", "H3": "2-5 years"}


def to_rgb(h):
    from docx.shared import RGBColor
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def shade_cell(cell, hexcolor):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def docx_new(title, subtitle=None):
    """Start a house-style deliverable Word document. Returns the Document."""
    from docx import Document
    from docx.shared import Pt, Inches
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    sec = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.9))
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = to_rgb(CLR_NAVY)
    if subtitle:
        s = doc.add_paragraph()
        sr = s.add_run(subtitle)
        sr.italic = True
        sr.font.size = Pt(11.5)
        sr.font.color.rgb = to_rgb(CLR_GREY)
    return doc


def docx_heading(doc, text, level=1):
    from docx.shared import Pt
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(15)
        r.font.color.rgb = to_rgb(CLR_TEAL)
    else:
        r.font.size = Pt(12.5)
        r.font.color.rgb = to_rgb(CLR_NAVY)
    return p


def docx_body(doc, text, italic=False):
    from docx.shared import Pt
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(11)
    return p


def docx_bullets(doc, items):
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def docx_callout(doc, text, heading="Assumptions this rests on"):
    """Shaded amber callout box for the demo note / assumption ledger."""
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade_cell(c, CLR_AMBER_BG)
    p = c.paragraphs[0]
    r = p.add_run(heading)
    r.bold = True
    r.font.size = Pt(10.5)
    for line in str(text).split("\n"):
        if line.strip():
            cp = c.add_paragraph()
            cr = cp.add_run(line.strip())
            cr.font.size = Pt(9.5)
    return t


def docx_table(doc, headers, rows, colwidths=None):
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(9.5)
        shade_cell(hdr[i], CLR_NAVY)
        r.font.color.rgb = to_rgb("FFFFFF")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            rr = p.add_run("" if v is None else str(v))
            rr.font.size = Pt(9)
    if colwidths:
        for i, w in enumerate(colwidths):
            for cell in t.columns[i].cells:
                cell.width = Inches(w)
    return t


def docx_figure(doc, path, width=6.6, caption=None):
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = to_rgb(CLR_GREY)


def docx_footer(doc, text):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.color.rgb = to_rgb(CLR_GREY)


def md_to_docx(doc, md):
    """Render a subset of Markdown (#/##/### headings, - bullets, paragraphs;
    strips ** markers) into an existing Document. For branch docs authored as
    markdown (PFDD plan, survey design)."""
    import re
    for block in str(md).split("\n"):
        s = block.rstrip()
        if not s.strip():
            continue
        if s.startswith("### "):
            docx_heading(doc, s[4:].strip(), level=2)
        elif s.startswith("## "):
            docx_heading(doc, s[3:].strip(), level=1)
        elif s.startswith("# "):
            docx_heading(doc, s[2:].strip(), level=1)
        elif s.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(re.sub(r"\*\*", "", s.lstrip()[2:]),
                              style="List Bullet")
        else:
            docx_body(doc, re.sub(r"\*\*", "", s))


def build_capacity_model(nodes):
    """Validate a capacity->strategy dependency model; return {nodes, edges}.

    Each node is a dict with: id, label, kind (a key of NODE_KIND_COLORS),
    effort (H1/H2/H3), phase (P1/P2/P3), cost, builds_on (list of ids),
    unlocks (list of ids). Edges combine both directions. This is the layer
    that ties CAPACITY/PERSONNEL investments (e.g. a Scientific Director hire)
    to the scientific/therapeutic steps they UNLOCK downstream."""
    by = {n["id"]: n for n in nodes}
    edges = set()
    for n in nodes:
        for u in n.get("unlocks", []):
            if u not in by:
                raise ValueError("%s unlocks unknown id %s" % (n["id"], u))
            edges.add((n["id"], u))
        for b in n.get("builds_on", []):
            if b not in by:
                raise ValueError("%s builds_on unknown id %s" % (n["id"], b))
            edges.add((b, n["id"]))
    return {"nodes": nodes, "edges": sorted(edges)}


def render_unlock_map(model, path="capacity_unlock_map.png", title=None,
                      phase_labels=None):
    """Render the capacity->strategy unlock figure: phase columns, left->right
    arrows, capacity-origin unlock edges emphasized in bold. `model` is the
    dict from build_capacity_model."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Patch
    from matplotlib.lines import Line2D
    nodes = model["nodes"]
    by = {n["id"]: n for n in nodes}
    edges = model["edges"]
    if phase_labels is None:
        phase_labels = {"P1": "Now  (0-12 mo)", "P2": "Near  (1-2 yr)",
                        "P3": "Later  (2-5 yr)"}
    phase_x = {"P1": 0, "P2": 1, "P3": 2}
    order = ["capacity", "infrastructure", "advocacy", "science",
             "funding", "therapeutic"]
    col = {"P1": [], "P2": [], "P3": []}
    for n in nodes:
        col[n["phase"]].append(n)
    pos = {}
    for ph, items in col.items():
        items.sort(key=lambda n: order.index(n["kind"])
                   if n["kind"] in order else 99)
        k = len(items)
        for i, n in enumerate(items):
            pos[n["id"]] = (phase_x[ph], k - 1 - i)
    maxrow = max(len(v) for v in col.values())
    fig, ax = plt.subplots(figsize=(16, 9.5))
    for ph, x in phase_x.items():
        ax.axvspan(x - 0.46, x + 0.46,
                   color="#f4f4f8" if x % 2 == 0 else "#eef1f6", zorder=0)
        ax.text(x, maxrow + 0.15, phase_labels.get(ph, ph), ha="center",
                va="bottom", fontsize=15, fontweight="bold", color="#22223b")
    bw, bh = 0.86, 0.72
    for a, b in edges:
        if a not in pos or b not in pos:
            continue
        xa, ya = pos[a]
        xb, yb = pos[b]
        is_unlock = (by[a]["kind"] == "capacity" and by[b]["kind"] in
                     ("therapeutic", "science", "funding", "infrastructure"))
        ec = "#6A4C93" if is_unlock else "#9aa0aa"
        lw = 2.6 if is_unlock else 1.3
        rad = 0.18 if yb != ya else 0.0
        ax.add_patch(FancyArrowPatch(
            (xa + bw / 2 * 0.62, ya), (xb - bw / 2 * 0.62, yb),
            connectionstyle="arc3,rad=%s" % rad, arrowstyle="-|>",
            mutation_scale=15, lw=lw, color=ec,
            alpha=0.9 if is_unlock else 0.5, zorder=2))
    for n in nodes:
        x, y = pos[n["id"]]
        c = "#" + NODE_KIND_COLORS.get(n["kind"], "888888")
        ax.add_patch(FancyBboxPatch(
            (x - bw / 2, y - bh / 2), bw, bh,
            boxstyle="round,pad=0.02,rounding_size=0.06", linewidth=1.5,
            edgecolor=c, facecolor=c + "22", zorder=3))
        ax.text(x, y + bh / 2 - 0.12, "%s . %s" % (n["id"], n.get("cost", "")),
                ha="center", va="top", fontsize=8.5, fontweight="bold",
                color=c, zorder=4)
        ax.text(x, y - 0.04, n["label"], ha="center", va="center",
                fontsize=8.6, color="#22223b", zorder=4, linespacing=1.15)
    handles = [Patch(facecolor="#" + NODE_KIND_COLORS[k] + "22",
                     edgecolor="#" + NODE_KIND_COLORS[k], label=k.title())
               for k in order if k in NODE_KIND_COLORS]
    handles.append(Line2D([0], [0], color="#6A4C93", lw=2.6,
                          label="Capacity unlocks ->"))
    handles.append(Line2D([0], [0], color="#9aa0aa", lw=1.3,
                          label="Builds-on dependency"))
    ax.legend(handles=handles, loc="lower center", ncol=4, frameon=False,
              fontsize=10.5, bbox_to_anchor=(0.5, -0.09))
    ax.set_xlim(-0.6, 2.6)
    ax.set_ylim(-1.0, maxrow + 0.9)
    ax.axis("off")
    ax.set_title(title or "Capacity-building -> scientific & therapeutic "
                 "unlocks\nInvesting in people and infrastructure now opens "
                 "strategic doors later", fontsize=17, fontweight="bold",
                 pad=18)
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def render_roadmap_views(options, which="both", path_prefix="capacity_roadmap",
                         demo_note=True):
    """Render the SAME options two ways so the leader can pick the frame that
    fits their board. `options` is a list of dicts (or DataFrame records) with
    keys ID, Horizon (H1/H2/H3), Cost (Low/Med/High), Posture, Option.
    which: 'effort' (menu by cost/capacity), 'timeline' (0-12mo/1-2yr/2-5yr),
    or 'both'. Returns list of written paths."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    import textwrap
    if hasattr(options, "to_dict"):
        options = options.to_dict("records")

    def wrap(s, n):
        return "\n".join(textwrap.wrap(str(s), n))

    def post_c(p):
        return "#" + POSTURE_COLORS.get(p, "888888")
    foot = ("DEMO org - a draft to cut and reorder, not a prescription."
            if demo_note else "Options to review, not a prescription.")
    paths = []
    if which in ("effort", "both"):
        tiers = [("Low cost / low capacity", "Low"),
                 ("Medium cost / building capacity", "Med"),
                 ("Higher cost / trial-readiness", "High")]
        maxrows = max(sum(1 for o in options if o["Cost"] == c)
                      for _, c in tiers)
        h = 0.85 / max(maxrows, 1)
        fig, axes = plt.subplots(1, 3, figsize=(17, 9.2))
        for ax, (title, cost) in zip(axes, tiers):
            ax.axis("off")
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
            ax.text(0.5, 0.985, title, ha="center", va="top", fontsize=13,
                    fontweight="bold", color="white",
                    bbox=dict(boxstyle="round,pad=0.5", facecolor="#22223b",
                              edgecolor="none"))
            rows = [o for o in options if o["Cost"] == cost]
            y = 0.90
            for r in rows:
                c = post_c(r["Posture"])
                ax.add_patch(FancyBboxPatch(
                    (0.02, y - h + 0.012), 0.96, h - 0.024,
                    boxstyle="round,pad=0.006,rounding_size=0.02",
                    linewidth=1.4, edgecolor=c, facecolor=c + "14"))
                ax.text(0.05, y - 0.016, r["ID"], fontsize=10.5,
                        fontweight="bold", color=c, va="top")
                ax.text(0.05, y - 0.052, r["Posture"], fontsize=6.8,
                        color="white", va="top",
                        bbox=dict(boxstyle="round,pad=0.22", facecolor=c,
                                  edgecolor="none"))
                ax.text(0.05, y - 0.088, wrap(r["Option"], 56), fontsize=7.3,
                        color="#22223b", va="top", linespacing=1.1)
                y -= h
        fig.suptitle("Capacity roadmap - VIEW A: options by cost & capacity "
                     "(low-cost-first menu)", fontsize=16, fontweight="bold",
                     y=0.995)
        fig.text(0.5, 0.02, "Posture pill colours the approach.  " + foot,
                 ha="center", fontsize=9.5, color="#555")
        p = path_prefix + "_effort_menu.png"
        fig.savefig(p, dpi=190, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        paths.append(p)
    if which in ("timeline", "both"):
        phases = ["0-12 months", "1-2 years", "2-5 years"]
        for o in options:
            o["_phase"] = HORIZON_TO_PHASE.get(o["Horizon"], o["Horizon"])
        colmax = max(sum(1 for o in options if o["_phase"] == p)
                     for p in phases)
        fig, ax = plt.subplots(figsize=(16, 9))
        ax.set_xlim(0, 3)
        ax.set_ylim(-0.6, colmax + 0.4)
        ax.axis("off")
        for i, ph in enumerate(phases):
            ax.axvspan(i + 0.03, i + 0.97,
                       color="#f4f4f8" if i % 2 == 0 else "#eef1f6", zorder=0)
            ax.text(i + 0.5, colmax + 0.25, ph, ha="center", fontsize=15,
                    fontweight="bold", color="#22223b")
            rows = [o for o in options if o["_phase"] == ph]
            for j, r in enumerate(rows):
                yy = colmax - 1 - j
                c = post_c(r["Posture"])
                ax.add_patch(FancyBboxPatch(
                    (i + 0.06, yy - 0.02), 0.88, 0.82,
                    boxstyle="round,pad=0.01,rounding_size=0.04",
                    linewidth=1.4, edgecolor=c, facecolor=c + "14", zorder=3))
                ax.text(i + 0.10, yy + 0.72, "%s . %s cost" %
                        (r["ID"], r["Cost"]), fontsize=8.5, fontweight="bold",
                        color=c, va="top", zorder=4)
                ax.text(i + 0.10, yy + 0.52, wrap(r["Option"], 42),
                        fontsize=7.6, color="#22223b", va="top", zorder=4)
        ax.set_title("Capacity roadmap - VIEW B: options on a timeline "
                     "(what happens when)", fontsize=16, fontweight="bold",
                     pad=16)
        fig.text(0.5, 0.02, "Same options as View A, placed on a calendar.  "
                 + foot, ha="center", fontsize=9.5, color="#555")
        p = path_prefix + "_timeline.png"
        fig.savefig(p, dpi=190, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        paths.append(p)
    return paths
